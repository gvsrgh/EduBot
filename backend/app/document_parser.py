"""
Document Parser Module

Extracts text content from uploaded documents in multiple formats:
- TXT: Direct read (UTF-8)
- PDF: Text-based extraction via PyPDF2, with OCR fallback for scanned pages
- DOCX: Paragraph extraction via python-docx

OCR support uses Tesseract (pytesseract + pdf2image) for scanned PDFs.
If Tesseract/poppler is not installed, OCR gracefully degrades with a warning.
"""

import io
import logging
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# PDF text extraction
# ---------------------------------------------------------------------------

def _extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from a PDF file.

    Strategy:
      1. Try PyPDF2 text extraction on every page.
      2. If a page yields little/no text, attempt OCR via Tesseract.
      3. Combine all pages into a single string.

    Returns the extracted text.
    """
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise RuntimeError("PyPDF2 is required for PDF parsing. Install it with: pip install PyPDF2")

    reader = PdfReader(io.BytesIO(content))
    all_text: list[str] = []
    ocr_attempted = False

    for page_num, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()

        # If the page has meaningful text, keep it
        if len(page_text) > 30:
            all_text.append(page_text)
        else:
            # Possibly a scanned page – try OCR
            ocr_text = _ocr_pdf_page(content, page_num)
            if ocr_text:
                ocr_attempted = True
                all_text.append(ocr_text)
            elif page_text:
                # Fall back to whatever little text PyPDF2 got
                all_text.append(page_text)

    if ocr_attempted:
        logger.info("OCR was used for one or more scanned pages in the PDF.")

    return "\n\n".join(all_text)


def _ocr_pdf_page(pdf_bytes: bytes, page_number: int) -> str:
    """
    Run OCR on a single PDF page using Tesseract + pdf2image.

    Returns extracted text, or empty string if OCR tools are unavailable.
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError:
        logger.warning(
            "pdf2image or pytesseract not installed – skipping OCR for scanned page %d. "
            "Install with: pip install pdf2image pytesseract  (and install Tesseract + Poppler on the system).",
            page_number,
        )
        return ""

    try:
        images = convert_from_bytes(
            pdf_bytes,
            first_page=page_number,
            last_page=page_number,
            dpi=300,
        )
        if not images:
            return ""
        text = pytesseract.image_to_string(images[0], lang="eng")
        return text.strip()
    except Exception as e:
        logger.warning("OCR failed for page %d: %s", page_number, e)
        return ""


# ---------------------------------------------------------------------------
# DOCX text extraction
# ---------------------------------------------------------------------------

def _extract_text_from_docx(content: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX parsing. Install it with: pip install python-docx")

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


# ---------------------------------------------------------------------------
# TXT (passthrough)
# ---------------------------------------------------------------------------

def _extract_text_from_txt(content: bytes) -> str:
    """Decode a plain-text file (UTF-8 with fallback to latin-1)."""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

PARSERS = {
    ".txt": _extract_text_from_txt,
    ".pdf": _extract_text_from_pdf,
    ".docx": _extract_text_from_docx,
}


def extract_text(filename: str, content: bytes) -> Tuple[str, str]:
    """
    Extract text from a document based on its file extension.

    Args:
        filename: Original filename (used to determine format).
        content: Raw file bytes.

    Returns:
        A tuple of (extracted_text, output_filename) where output_filename
        always has a .txt extension (for storage in the knowledge base).

    Raises:
        ValueError: If the file extension is not supported.
        RuntimeError: If a required library is missing.
    """
    ext = Path(filename).suffix.lower()

    parser = PARSERS.get(ext)
    if parser is None:
        supported = ", ".join(sorted(PARSERS.keys()))
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {supported}")

    text = parser(content)

    if not text.strip():
        raise ValueError(
            f"No text could be extracted from '{filename}'. "
            "If this is a scanned PDF, ensure Tesseract and Poppler are installed for OCR."
        )

    # Output filename: replace extension with .txt
    stem = Path(filename).stem
    output_filename = f"{stem}.txt"

    return text, output_filename
